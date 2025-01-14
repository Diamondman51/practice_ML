import os
from docx import Document
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import FAISS, Chroma, InMemoryVectorStore
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.documents import Document as LangChainDocument
from dotenv import load_dotenv

load_dotenv()

# API kalit
# os.environ["GOOGLE_API_KEY"] = "AIzaSyAh2r1Ci8WkKfsv8TBh5o_iXNn8ixDE-Kc"


# print(ord('в'))

# docs file yo'li
docx_files = [os.path.join('fayllar', doc) for doc in os.listdir('fayllar') if doc.endswith('docx')]

# DOCX ma'lumotlarini raw_text o'zgaruvchisiga o'qib olish
raw_text = ''
documents = []
for file in docx_files:
    doc = Document(file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    documents.append(LangChainDocument(page_content=text, metadata={"source": file}))

for docx_file in docx_files:
    doc = Document(docx_file)

    for paragraph in doc.paragraphs:
        raw_text += paragraph.text + "\n"

# matnni tokenlarga bo'lib olamiz. Bi ChatGPT talabidan chetga chiqib ketmaslik uchun
text_splitter = CharacterTextSplitter(
    separator="\n",
    chunk_size=3000,
    chunk_overlap=1000,
    length_function=len,
)

text_splitter2 = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
    chunk_size=3000,
    chunk_overlap=200
)

texts = text_splitter.split_text(raw_text)
texts2 = text_splitter2.split_documents(documents=documents)



# embeddingni yuklab olish - bu machine learning modellari va algoritmlari osongina foydalanishi mumkin bo'lgan ma'lumotlarni maxsus formatidir.
# Embedding - bu matn bo'lagi semantik ma'nosining zichlashtirilgan ifodasidir
# embeddings = OpenAIEmbeddings()
embeddings2 = GoogleGenerativeAIEmbeddings(model='models/text-embedding-004')

vector_store = InMemoryVectorStore(embeddings2)
# FAISS (Facebook AI oʻxshashlik qidiruvi) bu kutubxona boʻlib, ishlab chiquvchilarga bir-biriga oʻxshash multimedia hujjatlarini tezda izlash imkonini beradi.
# docsearch = FAISS.from_texts(texts, embeddings)
docsearch2 = Chroma.from_documents(documents=texts2, embedding=embeddings2)
# docsearch2 = docsearch2.as_retriever()

# ret_chain = docsearch2.map()

# LangChain - bu til modellari asosida ishlaydigan ilovalarni ishlab chiqish uchun framework.
from langchain.chains.question_answering import load_qa_chain
from langchain_openai.llms import OpenAI
from langchain_google_genai import GoogleGenerativeAI

# OpenAI-llm
# Hujjatlar to'plami bo'yicha Savol Javob qilish uchun foydalanishingiz mumkin bo'lgan zanjirni yuklaydi
# chain = load_qa_chain(OpenAI(), chain_type="stuff")
chain = load_qa_chain(GoogleGenerativeAI(model='gemini-1.5-flash'), chain_type="stuff")


def get_answer(query):
    # https://www.pinecone.io/learn/what-is-similarity-search/
    # vektor ko'rishdagi raqamlarni solishtirish

    docs = docsearch2.similarity_search(query)
    print('\n\nDocs: ', docs)
    result = chain.invoke({"input_documents": docs, "question": query + '\nотвечай только на русском языке'}, return_only_outputs=True)
    # print("\n\n", result)
    return result['output_text']

while True:
    query = input('\n\nВопрос: ')
    res = get_answer(query)
    # # os.system('cls')
    print('\n\nОтвет: ', res)
